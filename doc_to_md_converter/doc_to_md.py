"""
doc_to_md.py - Convertidor de archivos .doc/.docx a Markdown
Requisitos:
    pip install python-docx mammoth
"""

import sys
import os
import re
import argparse
from pathlib import Path

try:
    from docx import Document
    from docx.oxml.ns import qn
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import docx.opc.constants
except ImportError:
    print("ERROR: Instala python-docx con:  pip install python-docx")
    sys.exit(1)

try:
    import mammoth
except ImportError:
    print("ERROR: Instala mammoth con:  pip install mammoth")
    sys.exit(1)


# ─────────────────────────────────────────────
# Helpers de formato inline
# ─────────────────────────────────────────────

def run_to_md(run):
    """Convierte un Run de Word a texto Markdown preservando bold/italic/código."""
    text = run.text
    if not text:
        return ""

    # Evitar envolver espacios puros con marcadores
    stripped = text.strip()
    if not stripped:
        return text

    # Código (fuente monospace)
    font_name = (run.font.name or "").lower()
    is_code = any(m in font_name for m in ("courier", "consolas", "mono", "lucida console"))

    if is_code:
        return f"`{text}`"

    # Bold + Italic
    if run.bold and run.italic:
        return f"***{text}***"
    if run.bold:
        return f"**{text}**"
    if run.italic:
        return f"*{text}*"

    # Tachado
    if run.font.strike:
        return f"~~{text}~~"

    return text


def paragraph_inline(para):
    """Une todos los runs de un párrafo como Markdown inline."""
    return "".join(run_to_md(r) for r in para.runs)


# ─────────────────────────────────────────────
# Detección de estilo
# ─────────────────────────────────────────────

HEADING_MAP = {
    "heading 1": "#",
    "heading 2": "##",
    "heading 3": "###",
    "heading 4": "####",
    "heading 5": "#####",
    "heading 6": "######",
    "title": "#",
    "subtitle": "##",
}


def get_style_name(para):
    try:
        return para.style.name.lower()
    except Exception:
        return ""


def is_list_paragraph(para):
    style = get_style_name(para)
    return "list" in style or para._p.find(qn("w:numPr")) is not None


def get_list_level(para):
    numPr = para._p.find(qn("w:numPr"))
    if numPr is not None:
        ilvl = numPr.find(qn("w:ilvl"))
        if ilvl is not None:
            return int(ilvl.get(qn("w:val"), 0))
    return 0


def is_ordered_list(para):
    style = get_style_name(para)
    return "list number" in style or "numbered" in style


# ─────────────────────────────────────────────
# Tablas
# ─────────────────────────────────────────────

def table_to_md(table):
    rows = table.rows
    if not rows:
        return ""

    lines = []
    for i, row in enumerate(rows):
        cells = [" ".join(c.text.split()) for c in row.cells]
        lines.append("| " + " | ".join(cells) + " |")
        if i == 0:
            # Separador de encabezado
            lines.append("| " + " | ".join(["---"] * len(cells)) + " |")

    return "\n".join(lines)


# ─────────────────────────────────────────────
# Imágenes (exporta a carpeta images/)
# ─────────────────────────────────────────────

def extract_images(doc, output_dir: Path):
    """Extrae imágenes del documento y devuelve un dict {rId: ruta_relativa}."""
    images = {}
    images_dir = output_dir / "images"
    images_dir.mkdir(exist_ok=True)

    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            img_part = rel.target_part
            ext = img_part.content_type.split("/")[-1]
            ext = ext.replace("jpeg", "jpg")
            filename = f"{rel.rId}.{ext}"
            dest = images_dir / filename
            dest.write_bytes(img_part.blob)
            images[rel.rId] = f"images/{filename}"

    return images


def para_image_md(para, images):
    """Si el párrafo contiene una imagen inline, devuelve la línea Markdown."""
    blips = para._p.findall(".//" + qn("a:blip"))
    if not blips:
        return None

    lines = []
    for blip in blips:
        rId = blip.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
        if rId and rId in images:
            alt = paragraph_inline(para).strip() or "imagen"
            lines.append(f"![{alt}]({images[rId]})")

    return "\n".join(lines) if lines else None


# ─────────────────────────────────────────────
# Convertidor principal
# ─────────────────────────────────────────────

def docx_to_markdown(docx_path: Path, output_path: Path):
    doc = Document(str(docx_path))
    output_dir = output_path.parent
    images = extract_images(doc, output_dir)

    md_lines = []
    prev_was_list = False

    # Propiedades del documento (metadatos opcionales)
    core = doc.core_properties
    if core.title:
        md_lines.append(f"# {core.title}\n")
    if core.author:
        md_lines.append(f"> Autor: {core.author}\n")
    if core.description:
        md_lines.append(f"> {core.description}\n")
    if core.title or core.author or core.description:
        md_lines.append("")

    all_elements = list(doc.element.body)

    # Iteramos sobre los elementos del body para preservar el orden
    # (párrafos Y tablas en orden)
    doc_body = doc.element.body
    para_iter = iter(doc.paragraphs)
    table_iter = iter(doc.tables)

    # Mapeamos cada elemento XML a su objeto python-docx
    para_map = {}
    table_map = {}
    for p in doc.paragraphs:
        para_map[id(p._p)] = p
    for t in doc.tables:
        table_map[id(t._tbl)] = t

    for child in doc_body:
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag

        # ── Párrafo ──────────────────────────────────────────
        if tag == "p":
            para = para_map.get(id(child))
            if para is None:
                continue

            style = get_style_name(para)
            text = paragraph_inline(para)
            raw_text = para.text.strip()

            # Imagen inline
            img_md = para_image_md(para, images)
            if img_md:
                if prev_was_list:
                    md_lines.append("")
                md_lines.append(img_md)
                md_lines.append("")
                prev_was_list = False
                continue

            # Encabezados
            for key, prefix in HEADING_MAP.items():
                if style == key or style.startswith(key):
                    if prev_was_list:
                        md_lines.append("")
                    md_lines.append(f"{prefix} {text.strip()}")
                    md_lines.append("")
                    prev_was_list = False
                    break
            else:
                # Listas
                if is_list_paragraph(para):
                    level = get_list_level(para)
                    indent = "  " * level
                    bullet = "1." if is_ordered_list(para) else "-"
                    md_lines.append(f"{indent}{bullet} {text.strip()}")
                    prev_was_list = True

                # Bloque de código (estilo "Code" o "Preformatted")
                elif "code" in style or "preformat" in style or "verbatim" in style:
                    if prev_was_list:
                        md_lines.append("")
                    if not md_lines or md_lines[-1] != "```":
                        md_lines.append("```")
                    md_lines.append(raw_text)
                    # Cierre se maneja al salir del bloque (ver post-proceso)
                    prev_was_list = False

                # Cita / Quote
                elif "quote" in style or "block" in style:
                    if prev_was_list:
                        md_lines.append("")
                    md_lines.append(f"> {text.strip()}")
                    md_lines.append("")
                    prev_was_list = False

                # Párrafo normal / vacío
                else:
                    if prev_was_list:
                        md_lines.append("")
                        prev_was_list = False
                    if raw_text:
                        md_lines.append(text.strip())
                        md_lines.append("")
                    else:
                        # Párrafo vacío → separación
                        if md_lines and md_lines[-1] != "":
                            md_lines.append("")

        # ── Tabla ────────────────────────────────────────────
        elif tag == "tbl":
            tbl = table_map.get(id(child))
            if tbl is None:
                continue
            if prev_was_list:
                md_lines.append("")
            md_lines.append(table_to_md(tbl))
            md_lines.append("")
            prev_was_list = False

    # Post-proceso: cerrar bloques de código abiertos
    result = _close_code_blocks(md_lines)

    # Eliminar líneas en blanco consecutivas (máx. 2)
    result = _collapse_blank_lines(result)

    output_path.write_text(result, encoding="utf-8")
    print(f"✅  Convertido: {output_path}")


def _close_code_blocks(lines):
    """Asegura que cada ``` de apertura tenga su cierre."""
    open_block = False
    out = []
    for line in lines:
        if line == "```":
            if not open_block:
                open_block = True
                out.append(line)
            else:
                # No duplicar aperturas; lo tratamos como contenido
                out.append(line)
        elif open_block and line == "":
            # Bloque de código termina en línea vacía
            out.append("```")
            out.append("")
            open_block = False
        else:
            out.append(line)
    if open_block:
        out.append("```")
    return "\n".join(out)


def _collapse_blank_lines(text):
    return re.sub(r"\n{3,}", "\n\n", text)


# ─────────────────────────────────────────────
# Fallback con mammoth (para .doc binario)
# ─────────────────────────────────────────────

def doc_to_md_mammoth(doc_path: Path, output_path: Path):
    """Usa mammoth como fallback para archivos .doc (formato binario antiguo)."""
    with open(doc_path, "rb") as f:
        result = mammoth.convert_to_markdown(f)

    md = result.value
    if result.messages:
        print("Advertencias de mammoth:")
        for msg in result.messages:
            print(f"  - {msg}")

    output_path.write_text(md, encoding="utf-8")
    print(f"✅  Convertido con mammoth: {output_path}")


# ─────────────────────────────────────────────
# Entrada CLI
# ─────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(
        description="Convierte archivos .doc/.docx a Markdown sin perder datos."
    )
    parser.add_argument("input", help="Ruta al archivo .doc o .docx")
    parser.add_argument(
        "-o", "--output",
        help="Ruta de salida .md (por defecto: mismo nombre que el input)",
    )
    parser.add_argument(
        "--mammoth",
        action="store_true",
        help="Forzar uso de mammoth (útil para .doc binario antiguo)",
    )
    args = parser.parse_args()

    input_path = Path(args.input).resolve()
    if not input_path.exists():
        print(f"ERROR: No se encontró el archivo: {input_path}")
        sys.exit(1)

    if args.output:
        output_path = Path(args.output).resolve()
    else:
        output_path = input_path.with_suffix(".md")

    ext = input_path.suffix.lower()

    if args.mammoth or ext == ".doc":
        doc_to_md_mammoth(input_path, output_path)
    elif ext == ".docx":
        docx_to_markdown(input_path, output_path)
    else:
        print(f"ERROR: Extensión no soportada '{ext}'. Usa .doc o .docx")
        sys.exit(1)


if __name__ == "__main__":
    main()
